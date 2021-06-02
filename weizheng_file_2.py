import os
import sys
import easygui as g
import shutil
from docx import Document

title1 = '小唐制作'
filename1 = '一种WHAT'  # 你的模板名称，全部一致


def copy_file():  # 复制模板文件夹
    source_path = os.path.abspath(r'E:\唐国毅\专利撰写\一种WHAT')
    target_path = os.path.abspath(r'E:\唐国毅\专利撰写\ing\一种WHAT')

    '''if not os.path.exists(target_path):
        # 如果目标路径不存在原文件夹的话就创建
        os.makedirs(target_path)

    if os.path.exists(source_path):
        # 如果目标路径存在原文件夹的话就先删除
        shutil.rmtree(target_path)'''
    shutil.copytree(source_path, target_path)


def path_esists():  # 路径判断
    temp_path = os.getcwd()
    if os.path.exists(temp_path + '\\一种WHAT') == True:
        path = temp_path
    else:
        input_path = g.enterbox('输入你工作文件夹的存放路径（请直接去文件管理器复制）', title='小唐制作')
        while os.path.exists(input_path + '\\一种WHAT') == False:
            input_path = g.enterbox('存放路径错误请检查（直接去文件管理器复制）', title='小唐制作')
        path = input_path
    print(path)
    return path


def rename(path):  # 文件以及文件夹重命名
    filename2 = g.enterbox('输入你分配的案件名称', title=title1)  # 分配的案件名称
    user_choice = g.choicebox(msg='确定运行重命名文件？', title=title1, choices=('yes', 'no'))
    if user_choice == 'no':
        g.msgbox(msg='按任意键退出', title=title1)
        sys.exit(0)
    elif user_choice == 'yes':

        path1 = path + '\\' + filename1
        path2 = path + '\\' + filename2
        path2_1 = path2 + '\\' + filename1
        path2_2 = path2 + '\\' + filename2
        path2_2_1 = path2_2 + '\\' + filename1
        path2_2_2 = path2_2 + '\\' + filename2
        os.rename(path1, path2)  # 重命名第一层文件夹

        os.rename(path2_1, path2_2)  # 重命名第二层文件夹

        os.rename(path2_2_1 + '——说明书附图.pdf', path2_2_2 + '——说明书附图.pdf')

        os.rename(path2_2_1 + '——摘要附图.pdf', path2_2_2 + '——摘要附图.pdf')  # 重命名两个pdf文件

        os.rename(path2 + '\\' + '图' + '\\' + filename1 + '.SLDPRT', path2 + '\\' + '图' + '\\' + filename2 + '.SLDPRT')
        # 重命名空的sw文件，用于初始保存

        os.rename(path2 + '\\' + '图' + '\\' + filename1 + '.DWG', path2 + '\\' + '图' + '\\' + '模板' + filename2 + '.DWG')
        # 重命名模板DWG文件，用于区分文件

        os.rename(path2 + '\\' + filename1 + '——申请文件.docx',
                  path2 + '\\' + '未'+ filename2 + '——申请文件.docx')  # 重命名模板主文件夹里的初始word文件
        g.msgbox(msg='运行完成', title=title1)
        return filename2


def rename_self(path):
    filename2 = g.enterbox('输入你分配的案件名称', title=title1)  # 分配的案件名称
    path1 = path + '\\' + filename1
    path2 = path + '\\' + filename2
    path2_1 = path2 + '\\' + filename1
    path2_2 = path2 + '\\' + filename2
    path2_2_1 = path2_2 + '\\' + filename1
    path2_2_2 = path2_2 + '\\' + filename2
    os.rename(path1, path2)  # 重命名第一层文件夹

    os.rename(path2_1, path2_2)  # 重命名第二层文件夹

    os.rename(path2_2_1 + '——说明书附图.pdf', path2_2_2 + '——说明书附图.pdf')

    os.rename(path2_2_1 + '——摘要附图.pdf', path2_2_2 + '——摘要附图.pdf')  # 重命名两个pdf文件

    os.rename(path2 + '\\' + '图' + '\\' + filename1 + '.SLDPRT', path2 + '\\' + '图' + '\\' + filename2 + '.SLDPRT')
    # 重命名空的sw文件，用于初始保存

    os.rename(path2 + '\\' + '图' + '\\' + filename1 + '.DWG', path2 + '\\' + '图' + '\\' + '模板' + filename2 + '.DWG')
    # 重命名模板DWG文件，用于区分文件

    os.rename(path2 + '\\' + filename1 + '——申请文件.docx',
              path2 + '\\' + '未'+ filename2 + '——申请文件.docx')  # 重命名模板主文件夹里的初始word文件
    #g.msgbox(msg='运行完成', title=title1)
    return filename2


def replace_docx(docx_name):  # 替换word里的文字
    replace_choice = g.choicebox(msg='确定运行替代word模板里的案件名字？', title=title1, choices=('yes', 'no'))
    # 用户确认
    if replace_choice == 'yes':  # 执行修改word
        docx_file = Document('E:\\唐国毅\\专利撰写\\ing\\' + docx_name + '\\' + '未'+ docx_name + '——申请文件.docx')  # 读取文件
        for docx_text in docx_file.paragraphs:
            for run in docx_text.runs:  # 扫描word里面的文字
                if "WHAT" in run.text:  # 替换文字
                    if docx_name[:2] == '一种':
                        run.text = run.text.replace('WHAT', docx_name[2:])  # 如果存在一种，则提取一种后面的文字
                    else:
                        run.text = run.text.replace('WHAT', docx_name)  # 替换文字
            # print(docx_text.text)
            docx_file.save('E:\\唐国毅\\专利撰写\\ing\\' + docx_name + '\\' + '未'+ docx_name + '——申请文件.docx')
        g.msgbox(msg='运行完成', title=title1)
    else:
        sys.exit(0)


def replace_docx_self(docx_name):
    docx_file = Document('E:\\唐国毅\\专利撰写\\ing\\' + docx_name + '\\' + '未'+docx_name + '——申请文件.docx')  # 读取文件
    for docx_text in docx_file.paragraphs:
        for run in docx_text.runs:  # 扫描word里面的文字
            if "WHAT" in run.text:  # 替换文字
                if docx_name[:2] == '一种':
                    run.text = run.text.replace('WHAT', docx_name[2:])  # 如果存在一种，则提取一种后面的文字
                else:
                    run.text = run.text.replace('WHAT', docx_name)  # 替换文字
        # print(docx_text.text)
        docx_file.save('E:\\唐国毅\\专利撰写\\ing\\' + docx_name + '\\' + '未'+ docx_name + '——申请文件.docx')
    g.msgbox(msg='运行完成', title=title1)


if __name__ == "__main__":
    # path_out = path_esists()#判断文件夹是否存在

    user_choice_model = g.choicebox(msg='确定运行重命名文件？', title=title1, choices=('省心模式，一键到底', '谨慎模式，我来做主'))
    if user_choice_model == '谨慎模式，我来做主':
        copy_file()
        path_out = 'E:\\唐国毅\\专利撰写\\ing\\'
        docx_name = rename(path_out)
        replace_docx(docx_name)
    elif user_choice_model == '省心模式，一键到底':
        copy_file()
        path_out = 'E:\\唐国毅\\专利撰写\\ing\\'
        docx_name = rename_self(path_out)
        replace_docx_self(docx_name)
