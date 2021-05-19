import docx as d
from docx import Document
docx_file = Document(r'C:\Users\Administrator\Desktop\1.docx')#读取文件
for docx_text in docx_file.paragraphs:
    for run in docx_text.runs:#扫描word里面的文字
        if "唐国毅" in run.text:
            run.text = run.text.replace('唐国毅', 'WHAT')#替换文字
    print (docx_text.text)
    docx_file.save(r'C:\Users\Administrator\Desktop\1.docx')

#with open(r'C:\Users\Administrator\Desktop\1.docx',mode='r',encoding='Unicode') as word:
'''    word = word.read()
    print (word)
    print(word.read())
def read_document():
    document=Document('2.docx')WHAT
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if "唐诗" in run.text:
                run.text=run.text.replace('唐诗','宋词')
    document.save('3.docx')'''